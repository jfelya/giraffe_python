from docx import Document
from docx.shared import Inches


def run_doc():

    run = str(input("Do you want to create a document? (y) or (n): ").lower())
    if run == "n":
        print("That's so sad :(")
    else:
        # Opens up a new document
        document = Document()
        # document.add_page_break()
        # you can save the newly added part of the document in a variable for further use
        paragraph1 = document.add_paragraph("Lorem ipsum dolor sit amet.")
        # You can use the reference saved before to edit that object, like adding something before it
        paragraph1.insert_paragraph_before("This comes before the first paragraph")
        paragraph1.add_run("bold").bold = True
        paragraph1.add_run("and some")
        paragraph1.add_run("italic").italic = True

        # document.add_heading("Heading level 1", level=1)
        heading_insert = str(input("Do you want to add a heading to the doc? (y) or (n): ").lower())
        if heading_insert == "y":
            heading_text = str(input("Insert a name for your heading: "))
            document.add_heading(heading_text, 0)

        document.add_paragraph("Intense Quote", style="Intense Quote")
        document.add_paragraph("First item in an unordered list", style="List Bullet")
        document.add_paragraph("First item in an ordered list", style="List Number")

        image_insert = str(input("Do you want to add an image to the doc? (y) or (n): ").lower())
        if image_insert == "y":
            document.add_picture("pic1.JPG", width=Inches(1.5))

        table = document.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Column 1'
        hdr_cells[1].text = 'Column 2'
        hdr_cells[2].text = 'Column 3'
        document.add_page_break()

        try:
            file_name = str(input("Insert a name for your file: ").lower())
            document.save(f"{file_name}.docx")
            print(f"Congratulations, you have created a new document called {file_name}.docx!")
        except:
            print("Something terrible happened")


run_doc()
