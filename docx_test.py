from docx import Document
from docx.shared import Inches


def run_doc():

    run = str(input("¿Quieres crear un documento? (s) ó (n): ").lower())
    if run == "n":
        print("No se creó ningún documento")
    else:
        # Opens up a new document
        document = Document()
        # you can save the newly added paragraph of the document in a variable for further use

        def insert_paragraph():
            paragraph_text = str(input("Texto para el párrafo: "))
            document.add_heading(paragraph_text)
            print("Se agregó un párrafo\n")

        # You can use the reference saved before to edit that object, like adding something before it

        # document.add_heading("Heading level 1", level=1)
        def insert_heading():
            heading_text = str(input("Texto para el encabezado: "))
            document.add_heading(heading_text, 0)
            print("Se agregó un encabezado\n")

        # document.add_paragraph("Intense Quote", style="Intense Quote")
        def insert_list():
            list_text = str(input("Texto para el elemento lista: "))
            ordered_insert = str(input("¿Lo quieres numerado (n) o desorganizado (d)?: ").lower())
            if ordered_insert == "d":
                document.add_paragraph(list_text, style="List Bullet")
                print("Se agregó una lista desorganizada\n")
            if ordered_insert == "n":
                document.add_paragraph(list_text, style="List Number")
                print("Se agregó una lista numerada\n")

        def insert_image():
            document.add_picture("pic1.JPG", width=Inches(1.5))
            print("Se agregó una imágen\n")

        options = ""
        while options == "":
            print("¿Qué quieres agregar?: ")
            opt = str(input("p: Párrafo \ne: Encabezado \nl: Lista \ni: Imágen \ng: Guardar el documento \n").lower())

            if opt == "p":
                insert_paragraph()
            if opt == "e":
                insert_heading()
            if opt == "l":
                insert_list()
            if opt == "i":
                insert_image()
            if opt == "g":
                options = "holi"

        # TABLES
        table = document.add_table(rows=2, cols=2)
        hdr_cells1 = table.rows[0].cells
        hdr_cells1[0].text = "1"
        hdr_cells1[1].text = "2"
        hdr_cells2 = table.rows[1].cells
        hdr_cells2[0].text = "3"
        hdr_cells2[1].text = "4"
        document.add_page_break()

        try:
            file_name = str(input("Nombre para el documento: ").lower())
            document.save(f"{file_name}.docx")
            print(f"Felicidades, has creado un documento llamado: \"{file_name}.docx\"!")
        except:
            print("Algo ha pasado")


run_doc()
